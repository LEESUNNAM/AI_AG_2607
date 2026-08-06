"""
General-purpose .docx document builder using python-docx.

Styling defaults:
    - Font: Malgun Gothic (맑은 고딕) for both Latin and Korean text
    - Title size: 40pt, Body size: 20pt (Word has no "px" unit; pt is used)
    - Main color: customizable per document, applied to titles/headings
      and table headers for a consistent look

Usage as a library:
    from create_docx import DocBuilder

    doc = DocBuilder(title="Report Title", main_color="2E86AB")
    doc.add_heading("Section 1", level=1)
    doc.add_paragraph("Some body text.")
    doc.add_bullet_list(["Point A", "Point B", "Point C"])
    doc.add_table(
        headers=["Name", "Score"],
        rows=[["Alice", "90"], ["Bob", "85"]],
    )
    doc.save("output/report.docx")

Usage from the command line (creates a demo document):
    python create_docx.py output/demo.docx --color 2E86AB --title "Sample Document"
"""

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

DEFAULT_FONT = "맑은 고딕"
DEFAULT_TITLE_SIZE = 40
DEFAULT_BODY_SIZE = 20
DEFAULT_MAIN_COLOR = "1F4E79"
HEADING_SIZES = {1: 28, 2: 24, 3: 20}


def _parse_color(color: str | RGBColor) -> RGBColor:
    if isinstance(color, RGBColor):
        return color
    return RGBColor.from_string(color.lstrip("#").upper())


def _apply_font(run, font_name: str, size_pt: float, color: RGBColor | None = None, bold: bool | None = None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    # python-docx's font.name only sets the Latin typeface; Korean text needs
    # the East Asian typeface set explicitly or Word falls back to its default font.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def _shade_cell(cell, hex_color: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._element.get_or_add_tcPr().append(shd)


def _add_hyperlink(paragraph, url: str, text: str, font_name: str, size_pt: float):
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run_element = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), font_name)
    rpr.append(rfonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rpr.append(sz)
    run_element.append(rpr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


class DocBuilder:
    def __init__(
        self,
        title: str | None = None,
        main_color: str = DEFAULT_MAIN_COLOR,
        font_name: str = DEFAULT_FONT,
        title_size: float = DEFAULT_TITLE_SIZE,
        body_size: float = DEFAULT_BODY_SIZE,
    ):
        self.document = Document()
        self.main_color = _parse_color(main_color)
        self.font_name = font_name
        self.title_size = title_size
        self.body_size = body_size
        self._set_default_style()
        if title:
            self.add_title(title)

    def _set_default_style(self):
        normal = self.document.styles["Normal"]
        normal.font.name = self.font_name
        normal.font.size = Pt(self.body_size)
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), self.font_name)

    def add_title(self, text: str):
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        _apply_font(run, self.font_name, self.title_size, color=self.main_color, bold=True)
        return self

    def add_subtitle(self, text: str, size: float | None = None, color: str | RGBColor | None = None):
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        _apply_font(
            run,
            self.font_name,
            size or HEADING_SIZES.get(2, self.body_size),
            color=_parse_color(color) if color else self.main_color,
            bold=False,
        )
        return self

    def add_metadata(self, lines: list[str], size: float = 12):
        """Small centered info lines under the title, e.g. author/date/audience."""
        for line in lines:
            paragraph = self.document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line)
            _apply_font(run, self.font_name, size, color=RGBColor(0x59, 0x59, 0x59))
        return self

    def add_divider(self, char: str = "=", length: int = 70):
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(char * length)
        _apply_font(run, self.font_name, 10, color=RGBColor(0xAA, 0xAA, 0xAA))
        return self

    def add_references(self, items: list[tuple[str, str]], heading: str = "참고자료"):
        """items: list of (source_title, url) pairs, rendered as a numbered,
        clickable list — matches the report's own font/size for consistency."""
        self.add_heading(heading, level=1)
        for i, (source_title, url) in enumerate(items, start=1):
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run(f"{i}) {source_title}")
            _apply_font(run, self.font_name, self.body_size * 0.7)
            link_paragraph = self.document.add_paragraph()
            _add_hyperlink(link_paragraph, url, url, self.font_name, self.body_size * 0.7)
        return self

    def add_heading(self, text: str, level: int = 1, color: str | RGBColor | None = None, size: float | None = None):
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(text)
        heading_color = _parse_color(color) if color else self.main_color
        heading_size = size or HEADING_SIZES.get(level, self.body_size)
        _apply_font(run, self.font_name, heading_size, color=heading_color, bold=True)
        return self

    def add_paragraph(
        self,
        text: str,
        size: float | None = None,
        color: str | RGBColor | None = None,
        bold: bool = False,
        alignment: WD_ALIGN_PARAGRAPH | None = None,
    ):
        paragraph = self.document.add_paragraph()
        if alignment is not None:
            paragraph.alignment = alignment
        run = paragraph.add_run(text)
        _apply_font(
            run,
            self.font_name,
            size or self.body_size,
            color=_parse_color(color) if color else None,
            bold=bold,
        )
        return self

    def add_bullet_list(self, items: list[str]):
        for item in items:
            paragraph = self.document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(item)
            _apply_font(run, self.font_name, self.body_size)
        return self

    def add_numbered_list(self, items: list[str]):
        for item in items:
            paragraph = self.document.add_paragraph(style="List Number")
            run = paragraph.add_run(item)
            _apply_font(run, self.font_name, self.body_size)
        return self

    def add_table(self, headers: list[str], rows: list[list[str]], header_color: str | RGBColor | None = None):
        header_fill = _parse_color(header_color) if header_color else self.main_color
        table = self.document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = ""
            run = header_cells[i].paragraphs[0].add_run(header)
            _apply_font(run, self.font_name, self.body_size, color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)
            _shade_cell(header_cells[i], "%02X%02X%02X" % (header_fill[0], header_fill[1], header_fill[2]))
        for row in rows:
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = ""
                run = cells[i].paragraphs[0].add_run(str(value))
                _apply_font(run, self.font_name, self.body_size)
        return self

    def add_page_break(self):
        self.document.add_page_break()
        return self

    def save(self, path: str):
        out_path = Path(path)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(out_path)
        return out_path


def _demo(path: str, title: str, main_color: str):
    doc = DocBuilder(title=title, main_color=main_color)
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This document was generated by create_docx.py.")
    doc.add_heading("Key Points", level=1)
    doc.add_bullet_list(["First point", "Second point", "Third point"])
    doc.add_heading("Data", level=1)
    doc.add_table(
        headers=["Item", "Value"],
        rows=[["Alpha", "1"], ["Beta", "2"], ["Gamma", "3"]],
    )
    saved_path = doc.save(path)
    print(f"Saved: {saved_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate a demo .docx file.")
    parser.add_argument("output", nargs="?", default="output/demo.docx")
    parser.add_argument("--title", default="Sample Document")
    parser.add_argument("--color", default=DEFAULT_MAIN_COLOR, help="Main color as a hex string, e.g. 2E86AB")
    args = parser.parse_args()
    _demo(args.output, args.title, args.color)
